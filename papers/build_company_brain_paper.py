from __future__ import annotations

import re
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
)


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "company-brain-operational-memory.md"
DOCX_OUT = ROOT / "company-brain-operational-memory.docx"
PDF_OUT = ROOT / "company-brain-operational-memory.pdf"


def read_markdown() -> list[str]:
    return SOURCE.read_text(encoding="utf-8").splitlines()


def set_cell_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_docx_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    set_cell_shading(p, "F4F6F9")
    for index, line in enumerate(text.splitlines()):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
        run.font.size = Pt(8.5)


def configure_docx_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "Company Brain Technical Report"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if header.runs:
        header.runs[0].font.size = Pt(9)
        header.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    footer = section.footer.paragraphs[0]
    footer.text = "Human-Approved Operational Memory for Enterprise AI Agents"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer.runs:
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor(120, 120, 120)


def add_docx_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)


def build_docx(lines: list[str]) -> None:
    doc = Document()
    configure_docx_styles(doc)

    in_code = False
    code_lines: list[str] = []
    started_body = False

    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            if in_code:
                add_docx_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.name = "Calibri"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
            run.font.size = Pt(22)
            started_body = True
            i += 1
            continue

        if started_body and i <= 4:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line.strip())
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(80, 80, 80)
            i += 1
            continue

        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
            i += 1
            continue

        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            i += 1
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            p.add_run(line[2:].strip())
            i += 1
            continue

        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            p.add_run(re.sub(r"^\d+\. ", "", line).strip())
            i += 1
            continue

        add_docx_paragraph(doc, line.strip())
        i += 1

    doc.save(DOCX_OUT)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            "PaperTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=19,
            leading=23,
            alignment=TA_CENTER,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            "PaperMeta",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=12,
            textColor=colors.HexColor("#555555"),
            alignment=TA_CENTER,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            "H1Custom",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=14,
            leading=17,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=14,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            "H2Custom",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11.5,
            leading=14,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=10,
            spaceAfter=5,
        )
    )
    styles.add(
        ParagraphStyle(
            "BodyCustom",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9.8,
            leading=13.2,
            alignment=TA_LEFT,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            "CodeCustom",
            parent=styles["Code"],
            fontName="Courier",
            fontSize=7.2,
            leading=9,
            leftIndent=8,
            rightIndent=8,
            backColor=colors.HexColor("#F4F6F9"),
            borderColor=colors.HexColor("#DADCE0"),
            borderWidth=0.25,
            borderPadding=5,
            spaceBefore=4,
            spaceAfter=8,
        )
    )
    return styles


def para(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(escape(text), style)


def build_pdf(lines: list[str]) -> None:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=LETTER,
        rightMargin=0.85 * inch,
        leftMargin=0.85 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title="Company Brain: Human-Approved Operational Memory for Enterprise AI Agents",
        author="Aditya Shinde",
    )

    story = []
    in_code = False
    code_lines: list[str] = []
    meta_count = 0

    for line in lines:
        if line.startswith("```"):
            if in_code:
                story.append(Preformatted("\n".join(code_lines), styles["CodeCustom"]))
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            continue
        if line.startswith("# "):
            story.append(Paragraph(escape(line[2:].strip()), styles["PaperTitle"]))
            continue
        if meta_count < 3 and not line.startswith("## "):
            story.append(Paragraph(escape(line.strip()), styles["PaperMeta"]))
            meta_count += 1
            if meta_count == 3:
                story.append(Spacer(1, 12))
            continue
        if line.startswith("## References"):
            story.append(PageBreak())
            story.append(Paragraph("References", styles["H1Custom"]))
            continue
        if line.startswith("## "):
            story.append(Paragraph(escape(line[3:].strip()), styles["H1Custom"]))
            continue
        if line.startswith("### "):
            story.append(Paragraph(escape(line[4:].strip()), styles["H2Custom"]))
            continue
        if line.startswith("- "):
            story.append(
                ListFlowable(
                    [ListItem(para(line[2:].strip(), styles["BodyCustom"]), leftIndent=10)],
                    bulletType="bullet",
                    start="circle",
                    leftIndent=18,
                )
            )
            continue
        if re.match(r"^\d+\. ", line):
            story.append(
                ListFlowable(
                    [ListItem(para(re.sub(r"^\d+\. ", "", line).strip(), styles["BodyCustom"]), leftIndent=10)],
                    bulletType="1",
                    leftIndent=18,
                )
            )
            continue
        story.append(para(line.strip(), styles["BodyCustom"]))

    def footer(canvas, document):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#777777"))
        canvas.drawCentredString(
            LETTER[0] / 2,
            0.38 * inch,
            f"Company Brain Technical Report | Page {document.page}",
        )
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def main() -> None:
    lines = read_markdown()
    build_docx(lines)
    build_pdf(lines)
    print(f"Wrote {DOCX_OUT}")
    print(f"Wrote {PDF_OUT}")


if __name__ == "__main__":
    main()

